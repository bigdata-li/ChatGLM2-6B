import os
from transformers import AutoModel, AutoTokenizer
import gradio as gr
import mdtex2html
from utils import load_model_on_gpus
import re
import torch
from docx import Document
import time
import PyPDF2
import requests
from bs4 import BeautifulSoup
import io

ARTICLE_PROMPT_FOLDER = 'article_summary_prompts'
REPORT_PROMPT_FOLDER = 'report_summary_prompts'

status_state = gr.State("Welcome to ChifuGPT!")

def log(message):
    print(f"++++ {message}")

# Load content of a prompt file
def load_prompt_content(folder_path, prompt_file_name):
    file_path = buildPromptFileFullPath(folder_path, prompt_file_name)
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()
        
def update_status(message):
    status_state.value = message
    status_label.update(value=status_state.value)

# Callback for article_prompt_selector and report_prompt_selector
def on_selector_change(evt: gr.SelectData, selector_type: str, content_folder: str):
    global status_state
    current_selected_item = evt.value
    current_selector = selector_type

    log(f"Selected {current_selected_item} in {current_selector}")

    content = load_prompt_content(content_folder, current_selected_item)
    log(f"Loaded content of {current_selected_item} in {current_selector}: {content}")

    update_status(f"加載了提示詞文件 {current_selected_item} of {current_selector}。可以選擇'保存或創建提示詞文件'或者編輯新文件名後點擊保存")
    
    return content, status_state.value, current_selected_item

def on_article_selector_change(evt: gr.SelectData):
    return on_selector_change(evt, 'Article', ARTICLE_PROMPT_FOLDER)

def on_report_selector_change(evt: gr.SelectData):
    return on_selector_change(evt, 'Report', REPORT_PROMPT_FOLDER)

def on_dropdown_change_article(selected_option):
    selectorType = 'Article'
    # Bug per https://github.com/gradio-app/gradio/issues/5515
    # update_status(f"Changed custom value: {selected_option} in {selectorType}")
    log(f"Changed custom value: {selected_option} in {selectorType}.")
    return status_state.value

def on_dropdown_change_report(selected_option):
    selectorType = 'Report'
    # update_status(f"Changed custom value: {selected_option} in {selectorType}")
    log(f"Changed custom value: {selected_option} in {selectorType}")
    return status_state.value

# Save content to a prompt file
def save_prompt_content(folder_path, prompt_file_name, content, overwriteExist):
    file_path = buildPromptFileFullPath(folder_path, prompt_file_name)
    fileExist = os.path.exists(file_path)
    if fileExist and not overwriteExist:
        update_status(f"提示詞文件 {prompt_file_name} of {folder_path} 已經存在，請勾選'覆蓋已存在文件'。")
        return False

    saveOrCreate = "保存" if fileExist else "創建"
    saveOrUpdateEnglish = "Saving" if fileExist else "Creating"

    log(f"{saveOrUpdateEnglish} prompt file {prompt_file_name} of {folder_path}...")
    
    try:
        with open(file_path, 'w', encoding='utf-8') as file:
            file.write(content)
        
        update_status(f"{saveOrCreate}提示詞文件 {prompt_file_name} of {folder_path} 成功")
        return True
    except IOError as e:
        log(f"{saveOrCreate}提示詞文件 {prompt_file_name} of {folder_path} 失敗。Error: {e}")
        update_status(f"IOError: Failed to write to {file_path}. Error: {e}")
        return False

def save_create_prompt_file_article(content, selected_item, overwriteExistFile, promptFileName):
    return save_create_prompt_file('Article', article_prompt_selector, selected_item, content, overwriteExistFile, promptFileName)

def save_create_prompt_file_report(content, selected_item, overwriteExistFile, promptFileName):
    return save_create_prompt_file('Report', report_prompt_selector, selected_item, content, overwriteExistFile, promptFileName)

# Callback for save button
def save_create_prompt_file(selector_type, selector, selected_item, content, overwriteExistFile, promptFileName):
    global status_state

    log(f"Before-try Saving prompt file {promptFileName} of {selector.label}... selected_item={selected_item}")

    try:
        if not selector or not selected_item:
            update_status(f"請先選擇{selector.label}類型。")
            return None, status_state.value
        
        if content == "":
            update_status(f"請先輸入{selector.label}內容。")
            return None, status_state.value
        
        createNewFile = selected_item != promptFileName        
        
        folder_path = ARTICLE_PROMPT_FOLDER if selector_type == 'Article' else REPORT_PROMPT_FOLDER

        log(f"Before saving prompt file {promptFileName} of {selector.label}... current_selected_item={selected_item} createNewFile={createNewFile} folder_path={folder_path}")

        # Save the content to the new prompt file
        if not save_prompt_content(folder_path, promptFileName, content, overwriteExistFile):
            return None, status_state.value

        if createNewFile:
            log(f"Refreshing {selector.label} prompt selector...")

            # current_selected_item = promptFileName

            if selector_type == 'Article':
                article_summary_prompts[:] = load_prompts_from_folder(ARTICLE_PROMPT_FOLDER)
                log(f"Updating article_prompt_selector with choices: {article_summary_prompts}")
                return gr.Dropdown.update(choices=article_summary_prompts), status_state.value
            else:
                report_summary_prompts[:] = load_prompts_from_folder(REPORT_PROMPT_FOLDER)
                log(f"Updating report_prompt_selector with choices: {report_summary_prompts}")
                return gr.Dropdown.update(choices=report_summary_prompts), status_state.value
        else:
            return None, status_state.value

    finally:
        log(f"Finished saving prompt file {promptFileName} of {selector.label}.")
        # return None, status_state.value

def process_uploaded_files(file_wrappers):
    """
    Convert uploaded files into text.
    """
    content_list = []

    # Check if a single file is uploaded or multiple files
    if not isinstance(file_wrappers, list):
        file_wrappers = [file_wrappers]

    for file_wrapper in file_wrappers:
        file_name = file_wrapper.name
        file_content = file_wrapper.read()

        if file_name.endswith(".doc") or file_name.endswith(".docx"):
            print (f"++++ Reading Word document {file_name}.")
            # doc = Document(io.BytesIO(file_content))
            doc = Document(file_name)
            content = "\n".join([para.text for para in doc.paragraphs])
        elif file_name.endswith(".pdf"):
            print(f"++++ Reading PDF document {file_name}.")
            with open(file_name, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content = "\n".join([reader.pages[i].extract_text() for i in range(len(reader.pages))])
        elif file_name.endswith(".html"):
            print (f"++++ Reading and processing HTML document {file_name}...")
            soup = BeautifulSoup(file_content, 'lxml')
            for script in soup(["script", "style"]):
                script.decompose()
            content = " ".join(soup.stripped_strings)
        else:
            print (f"++++ Reading plain text document {file_name}...")
            content = file_content.decode('utf-8')

        # Wrap content
        wrapped_content = f"\n\n====上傳文檔開始(file://{file_name})===={{\n{content}\n}}====上傳文檔結束(file://{file_name})====\n"
        content_list.append(wrapped_content)

    return "\n".join(content_list)



def download_and_clean_html(url):
    try:
        print(f"++++ Downloading content from {url}...")
        response = requests.get(url)
        response.raise_for_status()

        print(f"++++ Cleaning up content from {url}...")
        soup = BeautifulSoup(response.text, 'lxml')
        # Remove unnecessary tags
        for script in soup(["script", "style"]):
            script.decompose()
        # Extract the text and clean it up
        return " ".join(soup.stripped_strings)
    except requests.RequestException as e:
        print(f"++++ Error fetching content from {url}: {e}")
        return f"Error fetching content from {url}: {e}"

def generate_report(article_summary_prompt, report_summary_prompt, urls, chatbot, max_length, top_p, temperature, history, past_key_values):
    urls = urls.split("\n")
    article_summaries = []

    summary_count = 1

    # 1. and 2. Download each article's content, concatenate it with the chosen article summary prompt
    for url in urls:
        content = download_and_clean_html(url)
        if not content:
            print(f"Skipping {url} due to download error or empty content.")
            continue
        article_input = f"{buildPromptFileFullPath(ARTICLE_PROMPT_FOLDER, article_summary_prompt)} {content}"

        # Use the model to generate a summary
        summary_response, new_history, new_past_key_values = predict(
            article_input, chatbot, max_length, top_p, temperature, history, past_key_values
        )
        if summary_response:
            summary = summary_response[-1][1]  # Take the last response (the generated summary)
            
            # Prepend the string with an incremental count
            formatted_summary = f"\n文章摘要 {summary_count}:\n{summary}"
            article_summaries.append(formatted_summary)
            
            # Increment the counter
            summary_count += 1
            
            history = new_history  # Update the history for the next iteration
            past_key_values = new_past_key_values  # Update past_key_values for the next iteration

    # 3. Collect all summaries and concatenate them to the reportSummaryPrompt
    combined_input = f"{buildPromptFileFullPath(REPORT_PROMPT_FOLDER, report_summary_prompt)} {' '.join(article_summaries)}"

    # Generate the final report
    report_response, _, _ = predict(
        combined_input, chatbot, max_length, top_p, temperature, history, past_key_values
    )
    if report_response:
        report = report_response[-1][1]  # Take the last response (the generated report)
        return report
    return "Error generating report"

def load_prompts_from_folder(folder_path):
    files_in_folder = os.listdir(folder_path)
    print(f"Files in {folder_path}: {files_in_folder}")
    return [file[:-4] for file in files_in_folder if file.endswith('.txt')]

def buildPromptFileFullPath(folder_path, prompt_file_name):
    return os.path.join(folder_path, prompt_file_name + '.txt')

print(f"Current Working Directory: {os.getcwd()}")
article_summary_prompts = load_prompts_from_folder(ARTICLE_PROMPT_FOLDER)
report_summary_prompts = load_prompts_from_folder(REPORT_PROMPT_FOLDER)

print(f"Article Summary Prompts: {article_summary_prompts}")
print(f"Report Summary Prompts: {report_summary_prompts}")

if torch.cuda.is_available():
    model_path = "../chatglm2-6b"
    print (f"CUDA is available. Loading model from {model_path}")
    tokenizer = AutoTokenizer.from_pretrained(model_path, trust_remote_code=True)
    model = load_model_on_gpus(model_path, num_gpus=2)
else: 
    model_path = "THUDM/chatglm2-6b"
    print (f"CUDA is not available. Loading model from {model_path}")
    tokenizer = AutoTokenizer.from_pretrained(model_path, trust_remote_code=True)
    model = AutoModel.from_pretrained(model_path, trust_remote_code=True).half().to('mps')
    
# model = AutoModel.from_pretrained("THUDM/chatglm2-6b", trust_remote_code=True).half().to('mps')
# tokenizer = AutoTokenizer.from_pretrained("../chatglm2-6b", trust_remote_code=True)
# model = AutoModel.from_pretrained("../chatglm2-6b", trust_remote_code=True).cuda()
# 多显卡支持，使用下面两行代替上面一行，将num_gpus改为你实际的显卡数量
# from utils import load_model_on_gpus
# model = load_model_on_gpus("../chatglm2-6b", num_gpus=2)
model = model.eval()

"""Override Chatbot.postprocess"""


def postprocess(self, y):
    if y is None:
        return []
    for i, (message, response) in enumerate(y):
        y[i] = (
            None if message is None else mdtex2html.convert((message)),
            None if response is None else mdtex2html.convert(response),
        )
    return y


gr.Chatbot.postprocess = postprocess


def parse_text(text):
    """copy from https://github.com/GaiZhenbiao/ChuanhuChatGPT/"""
    lines = text.split("\n")
    lines = [line for line in lines if line != ""]
    count = 0
    for i, line in enumerate(lines):
        if "```" in line:
            count += 1
            items = line.split('`')
            if count % 2 == 1:
                lines[i] = f'<pre><code class="language-{items[-1]}">'
                lines[i] = f'<br></code></pre>'
        else:
            if i > 0:
                if count % 2 == 1:
                    line = line.replace("`", "\`")
                    line = line.replace("<", "&lt;")
                    line = line.replace(">", "&gt;")
                    line = line.replace(" ", "&nbsp;")
                    line = line.replace("*", "&ast;")
                    line = line.replace("_", "&lowbar;")
                    line = line.replace("-", "&#45;")
                    line = line.replace(".", "&#46;")
                    line = line.replace("!", "&#33;")
                    line = line.replace("(", "&#40;")
                    line = line.replace(")", "&#41;")
                    line = line.replace("$", "&#36;")
                lines[i] = "<br>"+line
    text = "".join(lines)
    return text

def find_urls(text):
    url_pattern = re.compile(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+')
    return url_pattern.findall(text)


def predict(input, uploaded_files, chatbot, max_length, top_p, temperature, history, past_key_values, downloadUrlContent=False):
    old_input_length = len(input)

    if uploaded_files:
        file_contents = process_uploaded_files(uploaded_files)
        input += "\n" + file_contents
        
    if downloadUrlContent:
        # 1. Find all URLs in the input text
        urls = find_urls(input)

        # 2. Download the content of each URL and replace it in the input text
        for url in urls:
            content = download_and_clean_html(url)
            print (f"++++ Downloaded content from {url}: {content}")

            input = input.replace(url, "\n\n====網頁內容開始(" + url + ")===={\n\n" + content + "\n\n}====網頁內容結束(" + url + ")====\n")

    chatbot.append((parse_text(input), ""))

    update_status(f"Predicting started. input_length={len(input)} old_input_length={old_input_length}")
    log(status_state.value)

    # Get the current time before the loop starts
    start_time = time.time()

    for response, history, past_key_values in model.stream_chat(tokenizer, input, history, past_key_values=past_key_values,
                                                                return_past_key_values=True,
                                                                max_length=max_length, top_p=top_p,
                                                                temperature=temperature):
        chatbot[-1] = (parse_text(input), parse_text(response))
        update_status(f"Predicting in progress. input_length={len(input)} response_length={len(response)}")
        yield chatbot, history, past_key_values, status_state.value

    # Calculate the time difference after the loop completes
    end_time = time.time()
    duration = end_time - start_time

    status_message = f"Prediction completed in {duration:.2f} seconds. input_length={len(input)} and response_length={len(response)}"
    log(status_message)
    update_status(status_message)
    yield chatbot, history, past_key_values, status_state.value


def reset_user_input():
    return gr.update(value='')


def reset_state():
    return [], [], None

# sampleInput = """設想你是富蘭克林基金公司的研究員，請將以下金融文章的網頁內容總結成600字以內的摘要，請務必簡潔準確客觀。每一篇網頁內容都包括在其“網頁內容開始”和“網頁內容結束”之間。

# https://www.cnbc.com/2023/09/05/bond-yield-jump-is-not-death-to-equities-bofas-savita-subramanian.html
# """
sampleInput = """假設你是富蘭克林基金的研究員。請從以下金融文章中摘取要點，並將其總結為不超過1000字的報告，以供我們的客戶參考。請確保報告簡潔、準確且客觀。回答請使用繁體中文。如有參考資料或鏈接，請在報告末尾列出。

當您看到“網頁內容開始”和“網頁內容結束”時，其間的內容為網頁文章；當您看到“上傳文檔開始”和“上傳文檔結束”時，其間的內容為上傳的文件內容。

https://www.cnbc.com/2023/09/05/bond-yield-jump-is-not-death-to-equities-bofas-savita-subramanian.html

"""

with gr.Blocks() as demo:
    gr.HTML("""<h1 align="center">ChifuGPT v0.1</h1>""")

    chatbot = gr.Chatbot()
    with gr.Row():
        with gr.Column(scale=2):
            with gr.Column(scale=4):
                user_input = gr.Textbox(show_label=False, placeholder="Input...", lines=12, container=False)
            with gr.Column(scale=1):
                file_uploads = gr.Files(label="上傳本地Word/PDF/HTML文檔以嵌入提示", file_count="multiple")
                downloadUrlContent = gr.Checkbox(label="下載並嵌入提示中的http鏈接文章", value=True)
                submitBtn = gr.Button("Submit", variant="primary")
            
        with gr.Column(scale=2): # Adjust the 'scale' as per your design
            with gr.Row():
                article_prompt_selector = gr.Dropdown(choices=article_summary_prompts, label="Article總結提示詞") #, allow_custom_value=True)
                with gr.Row(flex=True):
                    with gr.Row():
                        promptFileNameArticle = gr.Textbox(show_label=False, lines=1, container=False)
                    with gr.Row():
                        save_article_button = gr.Button("保存或創建", size="sm", scale=0.5, min_width=100)
                        delete_article_prompt = gr.Button("刪除", size="sm", scale=0.5, min_width=100)

            with gr.Row():
                report_prompt_selector = gr.Dropdown(choices=report_summary_prompts, label="Report生成提示詞") #, allow_custom_value=True)
                with gr.Row(flex=True):
                    with gr.Row():
                        promptFileNameReport = gr.Textbox(show_label=False, lines=1, container=False)
                    with gr.Row():
                        save_report_button = gr.Button("保存或創建", size="sm", scale=0.5, min_width=100)
                        delete_report_prompt = gr.Button("刪除", size="sm", scale=0.5, min_width=100)

            with gr.Row():
                overwriteExistPromptFile = gr.Checkbox(label="覆蓋已存在文件", value=False)
            
            with gr.Row():
                with gr.Column(flex=1):
                    article_urls = gr.Textbox(label="引用文章鏈接列表", lines=6, placeholder="Enter URLs (one per line)", container=False, flex=1)
                    generate_report_btn = gr.Button("Generate Report")
                    status_label = gr.Textbox(label="Status", interactive=False, value=status_state.value)
                    
        with gr.Column(scale=1):
            emptyBtn = gr.Button("Clear History")
            max_length = gr.Slider(0, 32768, value=8192, step=1.0, label="Maximum length", interactive=True)
            top_p = gr.Slider(0, 1, value=0.8, step=0.01, label="Top P", interactive=True)
            temperature = gr.Slider(0, 1, value=0.95, step=0.01, label="Temperature", interactive=True)

    user_input.value = sampleInput
    
    history = gr.State([])
    past_key_values = gr.State(None)

    save_article_button.click(save_create_prompt_file_article, [user_input, article_prompt_selector, overwriteExistPromptFile, promptFileNameArticle], 
                              [article_prompt_selector, status_label], show_progress=True)
    save_report_button.click(save_create_prompt_file_report, [user_input, report_prompt_selector, overwriteExistPromptFile, promptFileNameReport],
                             [report_prompt_selector, status_label], show_progress=True)
    # create_prompt_button.click(create_new_prompt, [user_input], [status_label], show_progress=True)

    article_prompt_selector.select(on_article_selector_change, None, [user_input, status_label, promptFileNameArticle])
    report_prompt_selector.select(on_report_selector_change, None, [user_input, status_label, promptFileNameReport])

    # article_prompt_selector.change(on_dropdown_change_article, article_prompt_selector, status_label)
    # report_prompt_selector.change(on_dropdown_change_report, article_prompt_selector, status_label)

    submitBtn.click(predict, [user_input, file_uploads, chatbot, max_length, top_p, temperature, history, past_key_values, downloadUrlContent],
                    [chatbot, history, past_key_values, status_label], show_progress=True)
    submitBtn.click(reset_user_input, [], [user_input])

    emptyBtn.click(reset_state, outputs=[chatbot, history, past_key_values], show_progress=True)

    generate_report_btn.click(
        generate_report,
        [article_prompt_selector, report_prompt_selector, article_urls, chatbot, max_length, top_p, temperature, history, past_key_values],
        outputs=[gr.Textbox(label="Generated Report")]
    )

if torch.cuda.is_available():
    listen_ip = "0.0.0.0"
    listen_port = 7860
else:
    listen_ip = "127.0.0.1"
    listen_port = 8081

demo.queue().launch(server_name=listen_ip, server_port=listen_port, share=False, inbrowser=False)
